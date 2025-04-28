import os
import json
from typing import List, Dict, Optional
from langchain.schema.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
import PyPDF2
import docx

def load_knowledge_base(file_path: str) -> List[Document]:
    """Load knowledge base from JSON file"""
    with open(file_path, 'r', encoding='utf-8') as file:
        knowledge_base = json.load(file)
    
    documents = []
    for topic, content in knowledge_base.items():
        doc_text = f"# {topic}\n\n{content}"
        documents.append(Document(
            page_content=doc_text,
            metadata={"source": "knowledge_base", "topic": topic}
        ))
    
    return documents

def extract_text_from_pdf(pdf_path: str) -> str:
    """Extract text from PDF file"""
    text = ""
    with open(pdf_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(docx_path: str) -> str:
    """Extract text from DOCX file"""
    doc = docx.Document(docx_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def load_documents_from_directory(directory_path: str) -> List[Document]:
    """Load documents from directory"""
    documents = []
    for filename in os.listdir(directory_path):
        file_path = os.path.join(directory_path, filename)
        
        if filename.endswith('.pdf'):
            text = extract_text_from_pdf(file_path)
            documents.append(Document(
                page_content=text,
                metadata={"source": file_path, "filetype": "pdf"}
            ))
        
        elif filename.endswith('.docx'):
            text = extract_text_from_docx(file_path)
            documents.append(Document(
                page_content=text,
                metadata={"source": file_path, "filetype": "docx"}
            ))
        
        elif filename.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            documents.append(Document(
                page_content=text,
                metadata={"source": file_path, "filetype": "txt"}
            ))
    
    return documents

def split_documents(documents: List[Document], chunk_size: int = 1000, chunk_overlap: int = 100) -> List[Document]:
    """Split documents into chunks for better retrieval"""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ".", " ", ""]
    )
    return text_splitter.split_documents(documents)